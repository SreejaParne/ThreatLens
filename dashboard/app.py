from flask import Flask,send_file, render_template, request, redirect, url_for
from detector.attack_analyzer import generate_attack_graph
import os
import csv
import matplotlib
matplotlib.use("Agg")
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from datetime import datetime

app = Flask(__name__)

# ================= FILE PATHS =================
LOG_FILE = "logs/system.log"
ALERT_FILE = "logs/alerts.log"
EVENT_FILE = "logs/events.log"
ACK_FILE = "logs/acknowledged.log"
RESOLVED_FILE = "logs/resolved.log"
BLOCK_FILE = "logs/blocked_ips.log"

# ================= CLEAR FILES ON START =================
FILES_TO_CLEAR = [
    LOG_FILE,
    ALERT_FILE,
    EVENT_FILE,
    ACK_FILE,
    RESOLVED_FILE,
    BLOCK_FILE
]

for file in FILES_TO_CLEAR:
    os.makedirs(os.path.dirname(file), exist_ok=True)
    open(file, "w").close()

# ================= DASHBOARD =================
@app.route("/")
def dashboard():

    logs = open(LOG_FILE).readlines() if os.path.exists(LOG_FILE) else []
    alerts = open(ALERT_FILE).readlines() if os.path.exists(ALERT_FILE) else []
    acked = open(ACK_FILE).readlines() if os.path.exists(ACK_FILE) else []
    resolved = open("logs/resolved.log").readlines() if os.path.exists("logs/resolved.log") else []
    blocked = open("logs/blocked_ips.txt").readlines() if os.path.exists("logs/blocked_ips.txt") else []

    logs_count = len(logs)
    events_count = len(open(EVENT_FILE).readlines()) if os.path.exists(EVENT_FILE) else 0

    # clean lines
    alerts = [a.strip() for a in alerts]
    acked = [a.strip() for a in acked]
    resolved = [a.strip() for a in resolved]

    # blocked IP parsing (FIX)
    blocked = [b.strip() for b in blocked if "|" in b]

    blocked_ips = set()
    for line in blocked:
        line=line.strip()
        if "|" in line:
            parts = line.split("|")
            ip = parts[2].split("=")[1].strip()
            blocked_ips.add(ip)

    blocked_count = len(blocked_ips)

    # alert counts
    active_alerts = len([a for a in alerts if a not in acked and a not in resolved])
    ack_alerts = len([a for a in alerts if a in acked and a not in resolved])
    resolved_alerts = len(resolved)

    return render_template(
        "index.html",
        log_count=logs_count,
        event_count=events_count,
        active_alerts=active_alerts,
        ack_alerts=ack_alerts,
        resolved_alerts=resolved_alerts,
        blocked_count=blocked_count
    )


# ================= LOGS =================
@app.route("/logs")
def logs():
    logs = []

    if os.path.exists(LOG_FILE):
        with open(LOG_FILE) as f:
            for line in f.readlines()[-50:]:
                logs.append({
                    "message": line.strip(),
                    "source": "SSH" if "SSH" in line else "HTTP" if "HTTP" in line else "SYSTEM"
                })

    return render_template("logs.html", logs=logs)

# ================= ALERTS =================
@app.route("/alerts")
def alerts():
    alerts = []
    acknowledged = set()
    resolved = set()

    if os.path.exists(ACK_FILE):
        acknowledged = set(open(ACK_FILE).read().splitlines())

    if os.path.exists(RESOLVED_FILE):
        resolved = set(open(RESOLVED_FILE).read().splitlines())

    if os.path.exists(ALERT_FILE):
        alerts = open(ALERT_FILE).read().splitlines()

    unacked = [a for a in alerts if a not in acknowledged and a not in resolved]
    acked = [a for a in alerts if a in acknowledged and a not in resolved]
    resolved_alerts = [a for a in alerts if a in resolved]

    return render_template(
        "alerts.html",
        unacked_alerts=unacked,
        acked_alerts=acked,
        resolved_alerts=resolved_alerts
    )

# ================= ACK ALERT =================
@app.route("/ack", methods=["POST"])
def acknowledge():
    alert = request.form["alert"]

    with open(ACK_FILE, "a") as f:
        f.write(alert + "\n")

    return redirect(url_for("alerts"))

# ================= EVENTS =================
@app.route("/events")
def events():

    events_list = []

    if os.path.exists(EVENT_FILE):
        with open(EVENT_FILE) as f:
            for line in f.readlines():   # ⬅️ read ALL events safely
                try:
                    parts = [p.strip() for p in line.split("|")]

                    time = parts[0] if len(parts) > 0 else "N/A"
                    severity = parts[1] if len(parts) > 1 else "INFO"
                    message = parts[2] if len(parts) > 2 else "N/A"

                    ip = "N/A"
                    for p in parts:
                        if p.startswith("IP="):
                            ip = p.split("=", 1)[1]

                    events_list.append({
                        "time": time,
                        "severity": severity,
                        "message": f"{message} ({ip})"
                    })

                except Exception:
                    # SIEM must never crash on bad logs
                    continue

    return render_template("events.html", events=events_list)

# ================= STATS =================
@app.route("/stats")
def stats():
    generate_attack_graph()
    return render_template("stats.html")

# ================= REPORTS =================
@app.route("/reports")
def reports():
    return render_template("reports.html")

@app.route("/resolve", methods=["POST"])
def resolve():
    alert = request.form["alert"]

    with open(RESOLVED_FILE, "a") as f:
        f.write(alert + "\n")

    return redirect(url_for("alerts"))

@app.route("/blocked")
def blocked_ips():

    blocked = []

    if os.path.exists("logs/blocked_ips.txt"):
        with open("logs/blocked_ips.txt") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue

                # Expected format: time|status=BLOCKED|ip=1.2.3.4
                if "|" in line:
                    parts = line.split("|")
                    time = parts[0].strip()
                    ip = parts[2].split("=")[1].strip()

                    blocked.append({
                        "time": time,
                        "ip": ip
                    })

                # IGNORE old-format entries (no time info)
                else:
                    continue

    return render_template("blocked.html", blocked=blocked)


@app.route("/unblock/<ip>")
def unblock(ip):

    lines = open("logs/blocked_ips.txt").readlines()

    with open("logs/blocked_ips.txt", "w") as f:
        for line in lines:
            if ip not in line:
                f.write(line)

    return redirect("/blocked")

from flask import send_file
import os

@app.route("/download_report", methods=["POST"])
def download_report():
    report_format = request.form.get("format")

    # ===== LOAD DATA =====
    logs = open(LOG_FILE).readlines() if os.path.exists(LOG_FILE) else []
    events = open(EVENT_FILE).readlines() if os.path.exists(EVENT_FILE) else []
    alerts = open(ALERT_FILE).readlines() if os.path.exists(ALERT_FILE) else []
    acked = open(ACK_FILE).readlines() if os.path.exists(ACK_FILE) else []
    resolved = open(RESOLVED_FILE).readlines() if os.path.exists(RESOLVED_FILE) else []
    blocked = open("logs/blocked_ips.txt").readlines() if os.path.exists("logs/blocked_ips.txt") else []

    stats = {
        "Total Logs": len(logs),
        "Security Events": len(events),
        "Alerts": len(alerts),
        "Acknowledged": len(acked),
        "Resolved": len(resolved),
        "Blocked IPs": len(blocked)
    }

    os.makedirs("reports", exist_ok=True)

    # ================= CSV =================
    if report_format == "csv":
        path = "reports/threatlens_report.csv"
        with open(path, "w", newline="") as f:
            writer = csv.writer(f)

            writer.writerow(["ThreatLens Security Report"])
            writer.writerow(["Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
            writer.writerow([])

            writer.writerow(["=== STATISTICS ==="])
            for k, v in stats.items():
                writer.writerow([k, v])

            writer.writerow([])
            writer.writerow(["=== SYSTEM LOGS ==="])
            for l in logs:
                writer.writerow([l.strip()])

            writer.writerow([])
            writer.writerow(["=== SECURITY EVENTS ==="])
            for e in events:
                writer.writerow([e.strip()])

            writer.writerow([])
            writer.writerow(["=== ALERTS ==="])
            for a in alerts:
                writer.writerow([a.strip()])

            writer.writerow([])
            writer.writerow(["=== ACKNOWLEDGED ALERTS ==="])
            for a in acked:
                writer.writerow([a.strip()])

            writer.writerow([])
            writer.writerow(["=== RESOLVED ALERTS ==="])
            for r in resolved:
                writer.writerow([r.strip()])

            writer.writerow([])
            writer.writerow(["=== BLOCKED IPs ==="])
            for b in blocked:
                writer.writerow([b.strip()])

    # ================= PDF =================
    elif report_format == "pdf":
        path = "reports/threatlens_report.pdf"
        c = canvas.Canvas(path, pagesize=letter)

        y = 750
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, y, "ThreatLens Security Report")
        y -= 30

        c.setFont("Helvetica", 10)
        c.drawString(40, y, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        y -= 30

        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "Statistics Summary")
        y -= 20

        c.setFont("Helvetica", 10)
        for k, v in stats.items():
            c.drawString(60, y, f"{k}: {v}")
            y -= 15

        def section(title, lines):
            nonlocal y
            c.setFont("Helvetica-Bold", 12)
            y -= 20
            c.drawString(40, y, title)
            y -= 15
            c.setFont("Helvetica", 9)
            for line in lines:
                if y < 60:
                    c.showPage()
                    y = 750
                c.drawString(50, y, line.strip())
                y -= 12

        section("System Logs", logs)
        section("Security Events", events)
        section("Alerts", alerts)
        section("Acknowledged Alerts", acked)
        section("Resolved Alerts", resolved)
        section("Blocked IPs", blocked)

        c.save()

    # ================= DOC =================
    elif report_format == "doc":
        path = "reports/threatlens_report.docx"
        doc = Document()

        doc.add_heading("ThreatLens Security Report", level=1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading("Statistics Summary", level=2)
        for k, v in stats.items():
            doc.add_paragraph(f"{k}: {v}")

        def doc_section(title, lines):
            doc.add_heading(title, level=2)
            for l in lines:
                doc.add_paragraph(l.strip())

        doc_section("System Logs", logs)
        doc_section("Security Events", events)
        doc_section("Alerts", alerts)
        doc_section("Acknowledged Alerts", acked)
        doc_section("Resolved Alerts", resolved)
        doc_section("Blocked IPs", blocked)

        doc.save(path)

    else:
        return "Invalid format", 400

    return send_file(os.path.abspath(path), as_attachment=True)


# ================= RUN =================
if __name__ == "__main__":
    app.run(debug=True)