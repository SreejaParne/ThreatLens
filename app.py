import os
import random
from flask import Flask, send_file, render_template, request, redirect, url_for, session
import csv
import matplotlib
matplotlib.use("Agg")
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from datetime import datetime
from detector.attack_analyzer import generate_attack_graph
from fusion_engine.live_fusion import fetch_logs_by_time
from reports.forensic_report import generate_forensic_story
from datetime import timedelta
from real_time_sniffer import start_sniffer

# ---------------- APP INIT ----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")

app = Flask(__name__, template_folder=TEMPLATE_DIR)
app.secret_key = "threatlens_super_secret_key"
app.permanent_session_lifetime = timedelta(days=7)

from auth.auth_manager import AuthManager
from auth.rbac_engine import RBACEngine

auth = AuthManager()
rbac = RBACEngine()

# ---------------- FILE PATHS ----------------
LOG_DIR = os.path.join(BASE_DIR, "logs")
LOG_FILE = os.path.join(LOG_DIR, "system.log")
ALERT_FILE = os.path.join(LOG_DIR, "alerts.log")
EVENT_FILE = os.path.join(LOG_DIR, "events.log")
ACK_FILE = os.path.join(LOG_DIR, "acknowledged.log")
RESOLVED_FILE = os.path.join(LOG_DIR, "resolved.log")
BLOCK_FILE = os.path.join(LOG_DIR, "blocked_ips.txt")
POSITION_FILE = os.path.join(LOG_DIR, "engine_position.txt")

# ---------------- RESET LOGS ON EVERY RUN ----------------
def reset_logs():
    os.makedirs(LOG_DIR, exist_ok=True)
    for file in [
        LOG_FILE, ALERT_FILE, EVENT_FILE,
        ACK_FILE, RESOLVED_FILE, BLOCK_FILE, POSITION_FILE
    ]:
        open(file, "w").close()
    with open(POSITION_FILE, "w") as f:
        f.write("0")

# ---------------- AUTH HELPERS ----------------
def require_login():
    return "username" in session

def require_permission(permission):
    if not require_login():
        return redirect(url_for("login"))
    if not rbac.has_permission(session["role"], permission):
        return "Permission Denied"
    return None

# --------------REGISTER---------------
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        full_name = request.form["full_name"]
        org = request.form["organisation"]
        designation = request.form["designation"]
        email = request.form["email"]
        mobile = request.form["mobile"]
        role = request.form["role"]
        username = request.form["username"]
        password = request.form["password"]
        confirm = request.form["confirm_password"]

        if password != confirm:
            return render_template("register.html", error="Passwords do not match")

        success, msg = auth.register(full_name, org, designation, email, mobile, role, username, password)

        if success:
            return render_template("login.html", success="Registration successful! Please login.")

        return render_template("register.html", error=msg)

    return render_template("register.html")

# ---------------- LOGIN ----------------
@app.route("/login", methods=["GET", "POST"])
def login():
    # ❌ Do NOT auto-login if coming fresh after logout
    if request.method == "GET":
        return render_template("login.html")

    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        remember = request.form.get("remember")

        role = auth.authenticate(username, password)

        if role:
            session["username"] = username
            session["role"] = role

            if remember:
                session.permanent = True
            else:
                session.permanent = False

            return redirect(url_for("dashboard"))

        return render_template("login.html", error="Invalid credentials")

#--------------LOGOUT----------
@app.route("/logout")
def logout():
    session.clear()   # 🔥 removes everything (username, role, remember, etc.)
    return redirect(url_for("login"))

# -------------FORGOT PASSWORD------------
@app.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        username = request.form["username"]

        if username not in auth.users:
            return render_template("forgot_password.html", error="User not found")

        # 🔥 Generate OTP
        otp = str(random.randint(100000, 999999))

        session["reset_user"] = username
        session["otp"] = otp

        print("OTP for", username, ":", otp)  # 👉 demo purpose

        return redirect(url_for("verify_otp"))

    return render_template("forgot_password.html")

# ---------------VERIFY OTP-------------
@app.route("/verify_otp", methods=["GET", "POST"])
def verify_otp():
    if request.method == "POST":
        user_otp = request.form["otp"]
        new_password = request.form["password"]

        if user_otp == session.get("otp"):
            username = session.get("reset_user")

            auth.update_password(username, new_password)

            session.pop("otp", None)
            session.pop("reset_user", None)

            session["success_msg"] = "Password reset successful!"
            return redirect(url_for("login"))

        return render_template("verify_otp.html", error="Invalid OTP")

    return render_template("verify_otp.html")

# ---------------- DASHBOARD ----------------
@app.route("/")
def dashboard():
    if "username" not in session:
        return redirect(url_for("login"))

    # Load all files dynamically
    logs = open(LOG_FILE).readlines() if os.path.exists(LOG_FILE) else []
    alerts = open(ALERT_FILE).readlines() if os.path.exists(ALERT_FILE) else []
    acked = open(ACK_FILE).readlines() if os.path.exists(ACK_FILE) else []
    resolved = open(RESOLVED_FILE).readlines() if os.path.exists(RESOLVED_FILE) else []
    blocked_lines = open(BLOCK_FILE).readlines() if os.path.exists(BLOCK_FILE) else []

    blocked_ips = set()
    for line in blocked_lines:
        if "|" in line and "ip=" in line:
            ip = line.split("ip=")[1].strip()
            blocked_ips.add(ip)

    active_alerts = len([a for a in alerts if a not in acked and a not in resolved])
    ack_alerts = len([a for a in alerts if a in acked and a not in resolved])
    resolved_alerts = len(resolved)

    return render_template(
        "index.html",
        role=session["role"],
        log_count=len(logs),
        event_count=len(open(EVENT_FILE).readlines() if os.path.exists(EVENT_FILE) else []),
        active_alerts=active_alerts,
        ack_alerts=ack_alerts,
        resolved_alerts=resolved_alerts,
        blocked_count=len(blocked_ips)
    )
    
# -------------PROFILE------------
@app.route("/profile")
def profile():
    if not require_login():
        return redirect(url_for("login"))

    user = auth.get_user(session["username"])
    return render_template("profile.html", user=user, role=session["role"])

# ---------------UPDATE PROFILE---------------
@app.route("/edit_profile", methods=["GET", "POST"])
def edit_profile():
    if not require_login():
        return redirect(url_for("login"))

    user = auth.get_user(session["username"])

    if request.method == "POST":
        # Profile update
        full_name = request.form["full_name"]
        org = request.form["organisation"]
        designation = request.form["designation"]

        auth.update_user(session["username"], full_name, org, designation)

        # Password change (optional)
        old = request.form.get("old_password")
        new = request.form.get("new_password")
        confirm = request.form.get("confirm_password")

        if old and new and confirm:
            if new != confirm:
                return render_template("edit_profile.html", user=user, error="Passwords do not match")

            success, msg = auth.change_password(session["username"], old, new)

            if not success:
                return render_template("edit_profile.html", user=user, error=msg)

        return redirect(url_for("profile"))

    return render_template("edit_profile.html", user=user)

# -------------ADMIN VIEWING ALL USERS------------
@app.route("/admin/users")
def view_users():
    if not require_login():
        return redirect(url_for("login"))

    # Only admin allowed
    if session["role"] != "admin":
        return "Access Denied"

    users = auth.get_all_users()
    return render_template("admin_users.html", users=users)

@app.route("/delete_user/<username>")
def delete_user(username):
    if not require_login():
        return redirect(url_for("login"))

    if session["role"] != "admin":
        return "Access Denied"

    auth.delete_user(username)
    return redirect(url_for("view_users"))

# ---------------- LOGS ----------------
@app.route("/logs")
def logs_view():
    if not require_login():
        return redirect(url_for("login"))

    logs_data = []
    with open(LOG_FILE) as f:
        for line in f.readlines()[-50:]:
            line = line.strip()
            timestamp = line.split("|")[0] if "|" in line else ""
            logs_data.append({
                "timestamp": timestamp,
                "message": line,
                "source": "SSH" if "SSH" in line else "HTTP" if "HTTP" in line else "SYSTEM"
            })

    return render_template("logs.html", logs=logs_data, role=session["role"])

# ---------------- ALERTS ----------------
@app.route("/alerts")
def alerts_view():
    if not require_login():
        return redirect(url_for("login"))

    alerts = open(ALERT_FILE).read().splitlines() if os.path.exists(ALERT_FILE) else []
    acknowledged = set(open(ACK_FILE).read().splitlines() if os.path.exists(ACK_FILE) else [])
    resolved = set(open(RESOLVED_FILE).read().splitlines() if os.path.exists(RESOLVED_FILE) else [])

    unacked = [a for a in alerts if a not in acknowledged and a not in resolved]
    acked = [a for a in alerts if a in acknowledged and a not in resolved]
    resolved_alerts = [a for a in alerts if a in resolved]

    return render_template(
        "alerts.html",
        unacked_alerts=unacked,
        acked_alerts=acked,
        resolved_alerts=resolved_alerts,
        role=session["role"]
    )

@app.route("/ack", methods=["POST"])
def acknowledge():
    perm = require_permission("trigger_response")
    if perm:
        return perm
    alert = request.form["alert"]
    with open(ACK_FILE, "a") as f:
        f.write(alert + "\n")
    return redirect(url_for("alerts_view"))

@app.route("/resolve", methods=["POST"])
def resolve():
    perm = require_permission("trigger_response")
    if perm:
        return perm
    alert = request.form["alert"]
    with open(RESOLVED_FILE, "a") as f:
        f.write(alert + "\n")
    return redirect(url_for("alerts_view"))

# ---------------- SECURITY EVENTS ----------------
@app.route("/events")
def security_events():
    if not require_login():
        return redirect(url_for("login"))

    parsed_events = []
    if os.path.exists(EVENT_FILE):
        with open(EVENT_FILE, "r") as f:
            for line in f:
                parts = line.strip().split("|")
                if len(parts) >= 4:
                    parsed_events.append({
                        "time": parts[0],
                        "severity": parts[1],
                        "message": parts[2],
                        "ip": parts[3].split("=")[1] if "ip=" in parts[3] else "N/A"
                    })

    return render_template(
        "events.html",
        events=parsed_events,
        user=session["username"]
    )

# ---------------- BLOCKED IPS ----------------
@app.route("/blocked")
def blocked_ips_view():
    if not require_login():
        return redirect(url_for("login"))

    blocked = []
    if os.path.exists(BLOCK_FILE):
        with open(BLOCK_FILE) as f:
            for line in f:
                if "|" in line and "ip=" in line:
                    parts = line.split("|")
                    blocked.append({
                        "time": parts[0].strip(),
                        "ip": parts[2].split("=")[1].strip()
                    })

    return render_template("blocked.html", blocked=blocked, role=session["role"])

@app.route("/unblock/<ip>", methods=["POST"])
def unblock(ip):
    perm = require_permission("block_ip")
    if perm:
        return perm
    lines = open(BLOCK_FILE).readlines()
    with open(BLOCK_FILE, "w") as f:
        for line in lines:
            if ip not in line:
                f.write(line)
    return redirect(url_for("blocked_ips_view"))

# ---------------- STATS ----------------
from detector.statistics_engine import get_statistics

@app.route("/stats")
def stats():
    if not require_login():
        return redirect(url_for("login"))

    generate_attack_graph()   # THIS MUST RUN

    return render_template("stats.html", role=session["role"])
    
# ---------------- REPORTS ----------------
@app.route("/reports")
def reports_view():
    if not require_login():
        return redirect(url_for("login"))
    return render_template("reports.html", role=session["role"])

@app.route("/forensic_report")
def forensic_report_view():
    if not require_login():
        return redirect(url_for("login"))
    report_content = generate_forensic_story()
    return f"<pre>{report_content}</pre>"

@app.route("/download_report", methods=["POST"])
def download_report():
    perm = require_permission("export_reports")
    if perm:
        return perm

    report_format = request.form.get("format")
    logs = open(LOG_FILE).readlines() if os.path.exists(LOG_FILE) else []
    events = open(EVENT_FILE).readlines() if os.path.exists(EVENT_FILE) else []
    alerts = open(ALERT_FILE).readlines() if os.path.exists(ALERT_FILE) else []
    acked = open(ACK_FILE).readlines() if os.path.exists(ACK_FILE) else []
    resolved = open(RESOLVED_FILE).readlines() if os.path.exists(RESOLVED_FILE) else []
    blocked = open(BLOCK_FILE).readlines() if os.path.exists(BLOCK_FILE) else []

    forensic_story = generate_forensic_story()
    stats_data = {
        "Total Logs": len(logs),
        "Security Events": len(events),
        "Alerts": len(alerts),
        "Acknowledged": len(acked),
        "Resolved": len(resolved),
        "Blocked IPs": len(set(
            line.split("|")[2].split("=")[1].strip()
            for line in blocked if "|" in line
        ))
    }

    os.makedirs("reports", exist_ok=True)

    if report_format == "csv":
        path = "reports/threatlens_report.csv"
        with open(path, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["ThreatLens Security Report"])
            writer.writerow(["Generated", datetime.now()])
            writer.writerow([])
            writer.writerow(["Forensic Story"])
            writer.writerow([forensic_story])
            writer.writerow([])
            writer.writerow(["Statistics"])
            for k, v in stats_data.items():
                writer.writerow([k, v])
        return send_file(os.path.abspath(path), as_attachment=True)

    elif report_format == "pdf":
        path = "reports/threatlens_report.pdf"
        c = canvas.Canvas(path, pagesize=letter)
        y = 750
        c.drawString(40, y, "ThreatLens Security Report")
        y -= 30
        for line in forensic_story.split("\n"):
            c.drawString(50, y, line)
            y -= 15
        y -= 20
        for k, v in stats_data.items():
            c.drawString(50, y, f"{k}: {v}")
            y -= 15
        c.save()
        return send_file(os.path.abspath(path), as_attachment=True)

    elif report_format == "doc":
        path = "reports/threatlens_report.docx"
        doc = Document()
        doc.add_heading("ThreatLens Security Report", level=1)
        doc.add_paragraph(forensic_story)
        doc.add_heading("Statistics", level=2)
        for k, v in stats_data.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.save(path)
        return send_file(os.path.abspath(path), as_attachment=True)

    return "Invalid format", 400

# ---------------- INVESTIGATION ----------------
@app.route("/investigation", methods=["GET", "POST"])
def investigation_view():
    if not require_login():
        return redirect(url_for("login"))
    results = []
    if request.method == "POST":
        start = request.form.get("start_time")
        end = request.form.get("end_time")
        if start and end:
            start_time = datetime.strptime(start, "%Y-%m-%d %H:%M:%S")
            end_time = datetime.strptime(end, "%Y-%m-%d %H:%M:%S")
            results = fetch_logs_by_time(start_time, end_time)
    return render_template("investigation.html", results=results, role=session["role"])

# ---------------- RUN APP ----------------
if __name__ == "__main__":
    from threading import Thread
    from collector.log_collector import start_log_collector
    from detector.alert_engine import start_alert_engine

    # Reset logs on fresh run
    reset_logs()

    # Start background services
    Thread(target=start_log_collector, daemon=True).start()
    Thread(target=start_alert_engine, daemon=True).start()
    Thread(target=start_sniffer, daemon=True).start()

    # Run Flask
    app.run(host="localhost", port=9000, debug=False)